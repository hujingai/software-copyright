#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软著用户手册 DOCX 生成器
生成格式化的用户使用说明书 Word 文档，支持嵌入截图。
用法: python generate_manual_docx.py --info <采集信息JSON> --output <输出路径> [--screenshots <截图目录>]
依赖: pip install python-docx
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("[ERROR] 需要安装 python-docx: pip install python-docx")
    sys.exit(1)


# ============================================================
# 格式常量
# ============================================================
FONT_TITLE = '宋体'
FONT_BODY = '宋体'
TITLE_SIZE = Pt(16)
H1_SIZE = Pt(16)
H2_SIZE = Pt(14)
H3_SIZE = Pt(13)
BODY_SIZE = Pt(12)
SMALL_SIZE = Pt(10)
CAPTION_SIZE = Pt(9)
LINE_SPACING = 1.5

# A4 page
PAGE_WIDTH = Cm(21)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(2)


def set_run_font(run, font_name=FONT_BODY, size=BODY_SIZE, bold=False, color=None):
    """设置 Run 的中英文字体"""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color

    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_paragraph(doc, text="", font_name=FONT_BODY, size=BODY_SIZE, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_after=Pt(6),
                  first_line_indent=None, color=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = spacing_after
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent

    if text:
        run = p.add_run(text)
        set_run_font(run, font_name, size, bold, color)
    return p


def add_heading_paragraph(doc, text, level=1):
    """添加标题段落（使用自定义样式保持一致性）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)

    if level == 1:
        size = H1_SIZE
        prefix = ""
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:
        size = H2_SIZE
        prefix = ""
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        size = H3_SIZE
        prefix = ""
        alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(f"{prefix}{text}")
    set_run_font(run, FONT_TITLE, size, bold=True)
    return p


def add_section_number(doc, number: str, title: str):
    """添加带编号的章节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run(f"{number} {title}")
    set_run_font(run, FONT_TITLE, H1_SIZE, bold=True)
    return p


def add_image_with_caption(doc, image_path: str, caption_text: str, max_width=Cm(15)):
    """在文档中插入图片和题注"""
    if not image_path or not Path(image_path).exists():
        # 没有图片时插入占位文字
        p = add_paragraph(doc, f"【{caption_text}】", size=SMALL_SIZE,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(128, 128, 128))
        return

    try:
        # 插入图片
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(image_path, width=max_width)

        # 题注
        if caption_text:
            p = add_paragraph(doc, caption_text, size=CAPTION_SIZE,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            color=RGBColor(100, 100, 100))
    except Exception as e:
        print(f"  [WARN] 无法插入图片 {image_path}: {e}")
        p = add_paragraph(doc, f"【{caption_text} — 图片加载失败】", size=SMALL_SIZE,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(128, 128, 128))


def add_cover_page(doc, info: dict):
    """添加封面"""
    sw_name = info.get('software_full_name', '未命名软件')
    version = info.get('version', 'V1.0')
    owner = info.get('copyright_owner', '')

    # 空行
    for _ in range(6):
        add_paragraph(doc, "", spacing_after=Pt(0))

    # 软件名称
    add_paragraph(doc, sw_name, FONT_TITLE, Pt(22), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(16))

    # 手册标题
    add_paragraph(doc, "操作手册", FONT_TITLE, Pt(18), bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(24))

    # 版本信息
    add_paragraph(doc, f"版本号：{version}", FONT_BODY, Pt(14),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(8))

    if owner:
        add_paragraph(doc, f"著作权人：{owner}", FONT_BODY, Pt(14),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(8))

    add_paragraph(doc, f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}",
                  FONT_BODY, Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 分页
    doc.add_page_break()


def add_declaration_section(doc, info: dict):
    """添加声明章节"""
    sw_name = info.get('software_full_name', '本软件')

    add_section_number(doc, "一、", "引言")

    # 1.1 编写目的
    add_heading_paragraph(doc, "1.1 编写目的", level=2)
    add_paragraph(doc,
        f"本操作指南是{sw_name}的操作手册。开发此系统的目的是帮助用户实现"
        f"{info.get('business_domain', '业务管理')}的精准管理，提高工作效率，"
        f"增强数据可控性。通过数据化管理和自动化操作，使用户能够更加科学地完成工作任务。",
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        f"主要针对{info.get('target_users', '系统用户')}使用。"
        f"下面将以图文并茂的方式来说明。",
        first_line_indent=Cm(0.74))

    # 1.2 项目背景
    add_heading_paragraph(doc, "1.2 项目背景", level=2)
    add_paragraph(doc,
        info.get('software_overview',
                 f"{sw_name}是一款面向{info.get('target_users', '用户')}的"
                 f"{info.get('business_domain', '业务管理')}类应用软件。"),
        first_line_indent=Cm(0.74))

    # 1.3 术语说明
    glossary = info.get('glossary', '')
    if glossary:
        add_heading_paragraph(doc, "1.3 术语说明", level=2)
        for term in glossary.strip().split('\n'):
            term = term.strip()
            if term:
                add_paragraph(doc, f"● {term}", first_line_indent=Cm(0.74))

    doc.add_page_break()


def add_login_section(doc, info: dict, screenshot_dir: str = None):
    """添加登录章节"""
    sw_name = info.get('software_full_name', '本软件')

    add_section_number(doc, "二、", "登录方式及登录页面")

    add_heading_paragraph(doc, "2.1 登录页面概述", level=2)
    add_paragraph(doc,
        f"登录页面是用户进入{sw_name}的门户。在此页面，用户需要输入有效的用户名和密码，"
        f"以验证身份并获取系统的访问权限。",
        first_line_indent=Cm(0.74))

    add_heading_paragraph(doc, "2.2 打开登录页面", level=2)
    steps = [
        "1. 在您的网络浏览器中输入软件系统的网址。",
        "2. 按下回车键后，浏览器将显示登录页面。",
    ]
    for step in steps:
        add_paragraph(doc, step, first_line_indent=Cm(0.74))

    add_heading_paragraph(doc, "2.3 输入登录信息", level=2)
    login_steps = [
        '用户名/账号：在"用户名"或"账号"输入框中，键入您的用户名或账号。',
        '密码：在"密码"输入框中，键入您的密码。请注意，密码输入是隐藏的，以保护您的隐私。',
        '点击"登录"按钮，系统将验证您的身份信息。',
    ]
    for step in login_steps:
        add_paragraph(doc, step, first_line_indent=Cm(0.74))

    # 登录截图
    login_img = _find_screenshot(screenshot_dir, ['login', '登录', 'denglu'])
    add_image_with_caption(doc, login_img, "图：系统登录页面")

    add_heading_paragraph(doc, "2.4 安全提示", level=2)
    tips = [
        "确保在安全的环境下输入您的登录信息。",
        "不要在公共计算机或他人可见的情况下输入密码。",
        "定期更改您的密码，以维护账户安全。",
        "如连续多次输入错误密码，账户可能会被临时锁定。",
    ]
    for tip in tips:
        add_paragraph(doc, f"● {tip}", first_line_indent=Cm(0.74))

    doc.add_page_break()


def add_system_requirements_section(doc, info: dict):
    """添加系统要求章节"""
    add_section_number(doc, "三、", "系统要求")

    add_heading_paragraph(doc, "3.1 硬件环境", level=2)
    hardware = info.get('hardware_env',
        '处理器：Intel Core i5 及以上；内存：8GB 及以上；硬盘：256GB 及以上可用空间')
    add_paragraph(doc, hardware, first_line_indent=Cm(0.74))

    add_heading_paragraph(doc, "3.2 操作系统", level=2)
    os_env = info.get('os_env', 'Windows 10 及以上版本')
    add_paragraph(doc, os_env, first_line_indent=Cm(0.74))

    runtime = info.get('runtime_env', '')
    if runtime:
        add_heading_paragraph(doc, "3.3 运行支撑环境", level=2)
        add_paragraph(doc, runtime, first_line_indent=Cm(0.74))

    add_heading_paragraph(doc, "3.4 开发环境", level=2)
    dev_tools = info.get('dev_tools', 'VS Code / IntelliJ IDEA')
    prog_lang = info.get('programming_language', 'Java / Python')
    add_paragraph(doc, f"开发工具：{dev_tools}", first_line_indent=Cm(0.74))
    add_paragraph(doc, f"编程语言：{prog_lang}", first_line_indent=Cm(0.74))

    doc.add_page_break()


def add_features_section(doc, info: dict, screenshot_dir: str = None):
    """添加功能介绍及操作演示章节"""
    sw_name = info.get('software_full_name', '本软件')

    add_section_number(doc, "四、", "功能介绍及操作演示")

    modules = info.get('covered_modules', '')
    core_workflow = info.get('core_workflow', '')

    # 整体操作流程
    if core_workflow:
        add_paragraph(doc,
            f"{sw_name}的主要操作流程如下：{core_workflow}",
            first_line_indent=Cm(0.74))

    if not modules:
        add_paragraph(doc, "请根据软件实际功能模块填写各功能的操作说明。",
                     first_line_indent=Cm(0.74))
        return

    module_list = [m.strip() for m in modules.replace('、', '\n').replace(',', '\n').split('\n') if m.strip()]

    for idx, module in enumerate(module_list, 1):
        feature_num = f"4.{idx}"

        add_heading_paragraph(doc, f"{feature_num} {module}", level=2)

        # 功能概述
        add_paragraph(doc, f"{module}功能是{sw_name}的重要组成部分，"
                     f"旨在帮助用户高效完成{module}相关工作。"
                     f"该功能提供数据的新增、编辑、删除和查询等基本操作，"
                     f"用户可方便地对数据进行管理。"
                     f"系统界面简洁直观，操作流程清晰明了，"
                     f"大大提升了用户的工作效率。",
                     first_line_indent=Cm(0.74))

        # Main page screenshot
        main_img = _find_screenshot(screenshot_dir,
                                    [module.lower(), module.replace(' ', '').lower()])
        add_image_with_caption(doc, main_img, f"图：{module}主页面")

        # 新增操作
        add_heading_paragraph(doc, f"{feature_num}.1 新增{module}", level=3)
        add_paragraph(doc, f'1. 点击"新增"按钮，系统将跳转至{module}新增数据页面。',
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, "【必填字段】：根据页面提示，填写所有标记为必填的字段。",
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, "【选择字段】：对于需要选择值的字段，如下拉列表、单选框或多选框，根据需求进行选择。",
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, "【格式要求】：注意某些字段可能需要特定的格式，如日期、时间或数字格式。",
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, '2. 在确认所有信息正确无误后，点击"提交"按钮。',
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, "3. 系统将处理您的请求，并在数据添加成功后显示相应的成功信息。",
                     first_line_indent=Cm(0.74))

        add_img = _find_screenshot(screenshot_dir, [f'{module}_add', f'{module}_新增', f'{module}_new'])
        add_image_with_caption(doc, add_img, f"图：{module} — 新增数据页面")

        # 编辑操作
        add_heading_paragraph(doc, f"{feature_num}.2 编辑{module}", level=3)
        add_paragraph(doc,
            '1. 在数据列表页面中，找到需要编辑的数据记录。点击记录旁的【编辑】按钮，系统将显示编辑表单。',
            first_line_indent=Cm(0.74))
        add_paragraph(doc, "2. 编辑数据字段，根据需要，对数据表单中的字段进行编辑。",
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, "3. 审查编辑内容，在编辑完成后，仔细检查所有的变更，确保没有错误。",
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, '4. 保存编辑，在确认所有编辑无误后，点击"保存"按钮。系统将更新数据记录，并显示保存成功的提示。',
                     first_line_indent=Cm(0.74))

        edit_img = _find_screenshot(screenshot_dir, [f'{module}_edit', f'{module}_编辑'])
        add_image_with_caption(doc, edit_img, f"图：{module} — 编辑数据页面")

        # 删除操作
        add_heading_paragraph(doc, f"{feature_num}.3 删除{module}", level=3)
        add_paragraph(doc,
            "1. 单个删除：定位要删除的数据，在数据列表页面，浏览或搜索找到您想要删除的数据记录。",
            first_line_indent=Cm(0.74))
        add_paragraph(doc, '2. 选择数据记录，点击选中记录旁的"删除"按钮，系统将提示您确认删除操作。',
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, "3. 在确认删除的提示框中，仔细阅读提示信息，确保您理解删除操作的影响。",
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, '4. 点击"确认删除"后，数据将被永久移除。',
                     first_line_indent=Cm(0.74))

        delete_img = _find_screenshot(screenshot_dir, [f'{module}_delete', f'{module}_删除'])
        add_image_with_caption(doc, delete_img, f"图：{module} — 删除确认页面")

        # 每个模块之间留空
        if idx < len(module_list):
            add_paragraph(doc, "", spacing_after=Pt(12))


def add_faq_section(doc, info: dict):
    """添加常见问题章节"""
    sw_name = info.get('software_full_name', '本软件')

    add_section_number(doc, "五、", "常见问题 FAQ")

    faqs = [
        ("Q1：忘记登录密码怎么办？", "请联系系统管理员重置密码。如果您是管理员，请在服务器端使用密码重置工具进行重置。"),
        ("Q2：页面显示异常如何处理？", "请尝试刷新页面或清除浏览器缓存后重新登录。如问题依旧，请检查网络连接是否正常。"),
        ("Q3：数据保存失败是什么原因？", "请检查网络连接是否正常，确认必填字段是否都已填写完整。如仍无法解决，请联系技术支持。"),
        ("Q4：如何导出数据？", '在数据列表页面，点击"导出"按钮，选择导出格式（Excel/CSV/PDF），系统将生成导出文件。'),
        ("Q5：系统支持哪些浏览器？", "推荐使用 Chrome、Edge、Firefox 等主流浏览器的较新版本。"),
    ]

    for q, a in faqs:
        add_paragraph(doc, q, bold=True, spacing_after=Pt(4))
        add_paragraph(doc, a, first_line_indent=Cm(0.74), spacing_after=Pt(12))


def add_contact_section(doc, info: dict):
    """添加联系方式章节"""
    add_section_number(doc, "六、", "联系方式")

    add_paragraph(doc, "如在使用过程中遇到问题，请联系：", first_line_indent=Cm(0.74))
    add_paragraph(doc, f"著作权人：{info.get('copyright_owner', '')}",
                 first_line_indent=Cm(0.74))
    contact = info.get('contact_person', '')
    if contact:
        add_paragraph(doc, f"联系人：{contact}", first_line_indent=Cm(0.74))


def _find_screenshot(screenshot_dir: Optional[str], keywords: List[str]) -> Optional[str]:
    """在截图目录中查找匹配关键词的截图文件"""
    if not screenshot_dir:
        return None

    ss_dir = Path(screenshot_dir)
    if not ss_dir.exists() or not ss_dir.is_dir():
        return None

    # 支持的图片格式
    image_exts = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}

    # 收集所有图片
    all_images = [f for f in ss_dir.iterdir() if f.suffix.lower() in image_exts]

    if not all_images:
        return None

    # 按关键词匹配
    for kw in keywords:
        kw_lower = kw.lower().replace(' ', '')
        for img in all_images:
            name_lower = img.stem.lower().replace(' ', '').replace('_', '').replace('-', '')
            if kw_lower in name_lower:
                return str(img)

    # 如果没有精确匹配，返回第一个图片
    return str(all_images[0]) if all_images else None


def build_manual_docx(
    output_path: str,
    info: dict,
    screenshot_dir: Optional[str] = None,
):
    """构建完整的用户手册 DOCX 文档"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_BODY
    font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # 添加页眉
    sw_name = info.get('software_full_name', '未命名软件')
    version = info.get('version', 'V1.0')
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{sw_name}{version}　　操作手册")
    set_run_font(run, FONT_BODY, Pt(9))
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 封面
    add_cover_page(doc, info)

    # 引言
    add_declaration_section(doc, info)

    # 登录
    add_login_section(doc, info, screenshot_dir)

    # 系统要求
    add_system_requirements_section(doc, info)

    # 功能介绍
    add_features_section(doc, info, screenshot_dir)

    # FAQ
    add_faq_section(doc, info)

    # 联系方式
    add_contact_section(doc, info)

    # 保存
    doc.save(output_path)
    print(f"[OK] 用户手册 DOCX 已生成: {output_path}")
    return True


def load_info(json_path: str) -> dict:
    """加载采集信息"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def main():
    import argparse

    parser = argparse.ArgumentParser(description="生成软著用户手册 DOCX 文档")
    parser.add_argument("--info", required=True, help="采集信息 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 DOCX 文件路径")
    parser.add_argument("--screenshots", help="截图目录路径")
    parser.add_argument("--features", help="功能模块 JSON 文件路径（可选）")

    args = parser.parse_args()

    if not Path(args.info).exists():
        print(f"[ERROR] 采集信息文件不存在: {args.info}")
        sys.exit(1)

    info = load_info(args.info)
    screenshot_dir = args.screenshots

    if screenshot_dir and not Path(screenshot_dir).exists():
        print(f"[WARN] 截图目录不存在: {screenshot_dir}，将在手册中使用占位文字")
        screenshot_dir = None

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    build_manual_docx(str(output_path), info, screenshot_dir)


if __name__ == "__main__":
    main()
