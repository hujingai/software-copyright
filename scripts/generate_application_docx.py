#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软著申请表 DOCX 生成器
生成符合 CPCC 格式的计算机软件著作权登记申请表。
用法: python generate_application_docx.py --info <采集信息JSON> --output <输出路径>
依赖: pip install python-docx
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("[ERROR] 需要安装 python-docx: pip install python-docx")
    sys.exit(1)


# ============================================================
# 格式常量
# ============================================================
FONT_NAME = '宋体'
TITLE_SIZE = Pt(18)
HEADING_SIZE = Pt(14)
BODY_SIZE = Pt(11)
LABEL_SIZE = Pt(10)

# A4
PAGE_WIDTH = Cm(21)
PAGE_HEIGHT = Cm(29.7)
MARGIN = Cm(2.5)


def set_run_font(run, font_name=FONT_NAME, size=BODY_SIZE, bold=False):
    """设置 Run 的中英文字体"""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_table_row(table, cells, is_header=False):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

        if is_header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, FONT_NAME, BODY_SIZE, bold=True)
            # 设置单元格背景色
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)
        else:
            run = p.add_run(text)
            set_run_font(run, FONT_NAME, BODY_SIZE)


def add_form_table(doc, title: str, fields: List[tuple]):
    """添加表单表格"""
    # 章节标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, FONT_NAME, HEADING_SIZE, bold=True)

    # 创建表格
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    for row_data in fields:
        label, value = row_data
        add_table_row(table, [label, value])

    # 设置列宽（第一列约4cm，第二列自适应）
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.5)

    # 空行
    doc.add_paragraph()


def build_application_docx(output_path: str, info: dict):
    """构建申请表 DOCX"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("计算机软件著作权登记申请表")
    set_run_font(run, FONT_NAME, TITLE_SIZE, bold=True)

    sw_name = info.get('software_full_name', '')
    version = info.get('version', 'V1.0')
    owner = info.get('copyright_owner', '')

    # ========================
    # 一、软件基本信息
    # ========================
    software_fields = [
        ("软件全称", sw_name or "___________"),
        ("软件简称", info.get('software_short_name', '（无）')),
        ("版本号", version),
        ("软件分类", info.get('software_category', '应用软件')),
        ("软件说明", info.get('development_method', '原创')),
    ]
    add_form_table(doc, "一、软件基本信息", software_fields)

    # ========================
    # 二、开发信息
    # ========================
    dev_fields = [
        ("开发完成日期", info.get('completion_date', '')),
        ("发表状态", info.get('publication_status', '未发表')),
        ("首次发表日期", info.get('first_publication_date', '')),
        ("首次发表城市", info.get('first_publication_city', '')),
        ("开发目的", info.get('development_purpose', '提升工作效率，实现数据化管理')),
    ]
    add_form_table(doc, "二、开发与发表信息", dev_fields)

    # ========================
    # 三、软硬件环境
    # ========================
    env_fields = [
        ("开发硬件环境", info.get('dev_hardware_env', '通用PC（CPU 2.0GHz及以上、内存8GB及以上、硬盘256GB及以上）')),
        ("运行硬件环境", info.get('hardware_env', '处理器：Intel Core i5及以上；内存：8GB及以上；硬盘：256GB及以上可用空间')),
        ("开发操作系统", info.get('dev_os_env', 'Windows 10 / macOS 12')),
        ("运行操作系统", info.get('os_env', 'Windows 10 及以上版本')),
        ("开发工具", info.get('dev_tools', 'VS Code')),
        ("编程语言", info.get('programming_language', '')),
        ("运行支撑环境", info.get('runtime_env', '浏览器、HTTP 服务器')),
    ]
    add_form_table(doc, "三、软硬件环境", env_fields)

    # ========================
    # 四、功能与技术特点
    # ========================
    source_lines = info.get('source_lines', '')
    if not source_lines:
        source_lines = info.get('total_lines', '')

    tech_fields = [
        ("源程序量", str(source_lines) if source_lines else "___________"),
        ("面向领域/行业", info.get('business_domain', '')),
        ("主要功能",
         '\n'.join([m.strip() for m in (info.get('covered_modules', '') or '').replace('、', '\n').replace(',', '\n').split('\n') if m.strip()])
         or "___________"),
        ("技术特点",
         info.get('technical_features', '系统采用B/S架构，支持多用户并发访问；使用数据库进行数据持久化存储；界面简洁直观，操作流程清晰。')),
        ("技术特点选项",
         info.get('tech_category', '应用软件')),
    ]
    add_form_table(doc, "四、功能与技术特点", tech_fields)

    # ========================
    # 五、著作权人信息
    # ========================
    owner_fields = [
        ("姓名/名称", owner or "___________"),
        ("著作权人类型", info.get('copyright_owner_type', '个人')),
        ("证件号码", info.get('owner_id_number', '')),
        ("地址", info.get('owner_address', '')),
        ("联系电话", info.get('owner_phone', '')),
    ]
    add_form_table(doc, "五、著作权人信息", owner_fields)

    # ========================
    # 六、权利信息
    # ========================
    rights_fields = [
        ("权利取得方式", info.get('right_acquisition', '原始取得')),
        ("权利范围", info.get('right_scope', '全部权利')),
        ("开发方式", info.get('development_method', '独立开发')),
        ("权属补充说明", info.get('rights_note', '')),
    ]
    add_form_table(doc, "六、权利信息", rights_fields)

    # ========================
    # 七、申请人声明
    # ========================
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("七、申请人声明")
    set_run_font(run, FONT_NAME, HEADING_SIZE, bold=True)

    declaration_text = (
        "本人/本单位保证所提交的申请材料真实、准确、完整，"
        "所申请的软件为本人/本单位独立开发，未侵犯他人合法权益。"
        "如有不实之处，本人/本单位愿承担由此引起的法律责任。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(declaration_text)
    set_run_font(run, FONT_NAME, BODY_SIZE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("申请人签字/盖章：________________")
    set_run_font(run, FONT_NAME, BODY_SIZE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"申请日期：{datetime.now().strftime('%Y年%m月%d日')}")
    set_run_font(run, FONT_NAME, BODY_SIZE)

    # 保存
    doc.save(output_path)
    print(f"[OK] 申请表 DOCX 已生成: {output_path}")
    return True


def load_info(json_path: str) -> dict:
    """加载采集信息"""
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def main():
    import argparse

    parser = argparse.ArgumentParser(description="生成软著申请表 DOCX 文档")
    parser.add_argument("--info", required=True, help="采集信息 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 DOCX 文件路径")

    args = parser.parse_args()

    if not Path(args.info).exists():
        print(f"[ERROR] 采集信息文件不存在: {args.info}")
        sys.exit(1)

    info = load_info(args.info)
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    build_application_docx(str(output_path), info)


if __name__ == "__main__":
    main()
