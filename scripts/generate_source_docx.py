#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
软著源代码文档 DOCX 生成器
从源码分页数据生成格式化的 Word 文档（符合 CPCC 格式规范）。
用法: python generate_source_docx.py --input <分页JSON> --output <输出路径> --name <软件名称> --version <版本号>
依赖: pip install python-docx
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("[ERROR] 需要安装 python-docx: pip install python-docx")
    sys.exit(1)


# ============================================================
# 格式常量
# ============================================================
LINES_PER_PAGE = 50
CODE_FONT_NAME = 'Courier New'
CODE_FONT_SIZE = Pt(8)
CHINESE_FONT_NAME = '宋体'
CHINESE_FONT_SIZE = Pt(12)
HEADER_FONT_SIZE = Pt(9)
PAGE_MARGIN_TOP = Cm(2)
PAGE_MARGIN_BOTTOM = Cm(2)
PAGE_MARGIN_LEFT = Cm(3)
PAGE_MARGIN_RIGHT = Cm(2)

# A4 size in EMU
A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)


def set_cell_font(run, font_name=CHINESE_FONT_NAME, font_size=CHINESE_FONT_SIZE, bold=False):
    """设置字体属性"""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_code_font(run):
    """设置代码字体"""
    run.font.name = CODE_FONT_NAME
    run.font.size = CODE_FONT_SIZE
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CODE_FONT_NAME)
    rFonts.set(qn('w:ascii'), CODE_FONT_NAME)
    rFonts.set(qn('w:hAnsi'), CODE_FONT_NAME)


def add_header_to_section(section, sw_name: str, version: str, doc_type: str = "源代码文档"):
    """给节添加页眉"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{sw_name} {version}　　{doc_type}")
    set_cell_font(run, CHINESE_FONT_NAME, HEADER_FONT_SIZE)


def add_footer_to_section(section, sw_name: str, version: str):
    """给节添加页脚"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # 添加页码域
    run = p.add_run("第 ")
    set_cell_font(run, CHINESE_FONT_NAME, Pt(9))
    # 插入页码
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p.add_run()
    run1._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run1._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run1._element.append(fldChar2)
    run2 = p.add_run(" 页")
    set_cell_font(run2, CHINESE_FONT_NAME, Pt(9))


def create_cover_page(doc, sw_name: str, version: str, total_pages: int, strategy: str):
    """创建封面页"""
    # 空行
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # 软件名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(sw_name)
    set_cell_font(run, CHINESE_FONT_NAME, Pt(22), bold=True)

    doc.add_paragraph()  # 空行

    # 版本号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"版本号：{version}")
    set_cell_font(run, CHINESE_FONT_NAME, Pt(14))

    for _ in range(3):
        doc.add_paragraph()

    # 文档类型
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("源代码文档")
    set_cell_font(run, CHINESE_FONT_NAME, Pt(22), bold=True)

    for _ in range(2):
        doc.add_paragraph()

    # 信息表
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_items = [
        ("总页数", f"{total_pages} 页"),
        ("提交策略", "全部提交" if strategy == "submit_all" else "前30页 + 后30页"),
        ("每页行数", f"{LINES_PER_PAGE} 行"),
        ("生成日期", datetime.now().strftime('%Y年%m月%d日')),
    ]
    for i, (label, value) in enumerate(info_items):
        cell_label = info_table.cell(i, 0)
        cell_label.width = Cm(4)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label + "：")
        set_cell_font(run, CHINESE_FONT_NAME, Pt(11), bold=True)

        cell_value = info_table.cell(i, 1)
        cell_value.width = Cm(10)
        p = cell_value.paragraphs[0]
        run = p.add_run(value)
        set_cell_font(run, CHINESE_FONT_NAME, Pt(11))

    # 分页
    doc.add_page_break()


def add_code_page(doc, page_data: dict, sw_name: str, version: str):
    """
    添加一页代码
    page_data: {'page_number': int, 'source_file': str, 'lines': [str]}
    """
    # 页眉信息（作为段落）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{sw_name} {version}　　　　　　　　　　第 {page_data['page_number']} 页 / 共 Y 页")
    set_cell_font(run, CHINESE_FONT_NAME, Pt(8))
    # 设置浅色
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 代码行
    for line in page_data.get('lines', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)

        # 转义特殊字符
        safe_line = line.replace('\t', '    ')
        if not safe_line:
            safe_line = ' '

        run = p.add_run(safe_line[:130])  # 截断过长行
        set_code_font(run)

    # 文件来源（页脚）
    fname = Path(page_data['source_file']).name if page_data.get('source_file') else ''
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"文件：{fname}")
    set_cell_font(run, CHINESE_FONT_NAME, Pt(7))
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 页面分隔
    doc.add_page_break()


def build_source_docx(
    output_path: str,
    pages_data: dict,
    sw_name: str,
    version: str,
):
    """
    构建源代码 DOCX 文档
    pages_data 格式:
    {
        "first_N_pages": [{page_number, source_file, lines}, ...],
        "last_N_pages": [{page_number, source_file, lines}, ...],
        "total_pages": int,
        "strategy": "submit_all" | "first_30_last_30"
    }
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = PAGE_MARGIN_TOP
    section.bottom_margin = PAGE_MARGIN_BOTTOM
    section.left_margin = PAGE_MARGIN_LEFT
    section.right_margin = PAGE_MARGIN_RIGHT

    # 添加页眉页脚
    add_header_to_section(section, sw_name, version)
    add_footer_to_section(section, sw_name, version)

    # 封面
    create_cover_page(
        doc, sw_name, version,
        pages_data.get('total_pages', 0),
        pages_data.get('strategy', 'submit_all')
    )

    # 代码页
    all_pages = pages_data.get('all_pages', [])
    first_n = pages_data.get('first_N_pages', [])
    last_n = pages_data.get('last_N_pages', [])

    if all_pages and not first_n and not last_n:
        # submit_all 模式
        pages_to_render = all_pages
    else:
        pages_to_render = first_n + last_n

    # 如果有前30页和后30页，加一个分隔标记
    show_separator = bool(first_n and last_n)

    for i, page in enumerate(pages_to_render):
        add_code_page(doc, page, sw_name, version)

        # 在前30和后30之间加分隔页
        if show_separator and i == len(first_n) - 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(200)
            run = p.add_run("— — — 以下为后30页 — — —")
            set_cell_font(run, CHINESE_FONT_NAME, Pt(12))
            doc.add_page_break()

    # 保存
    doc.save(output_path)
    print(f"[OK] 源代码 DOCX 已生成: {output_path}")
    return True


def load_pages_from_json(json_path: str) -> dict:
    """从 JSON 文件加载分页数据"""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 支持两种 JSON 格式
    if 'pages' in data:
        # generate_source_doc.py 的输出格式
        result = {
            'total_pages': data.get('metadata', {}).get('total_pages', len(data['pages'])),
            'strategy': 'submit_all' if len(data['pages']) <= 60 else 'first_30_last_30',
        }
        pages = data['pages']
        if len(pages) <= 60:
            result['all_pages'] = pages
        else:
            result['first_N_pages'] = pages[:30]
            result['last_N_pages'] = pages[-30:]
        return result
    else:
        # 直接是分页结果
        return data


def main():
    import argparse

    parser = argparse.ArgumentParser(description="生成软著源代码 DOCX 文档")
    parser.add_argument("--input", required=True, help="分页数据 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 DOCX 文件路径")
    parser.add_argument("--name", required=True, help="软件全称")
    parser.add_argument("--version", default="V1.0", help="版本号")

    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"[ERROR] 输入文件不存在: {args.input}")
        sys.exit(1)

    pages_data = load_pages_from_json(args.input)
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    build_source_docx(str(output_path), pages_data, args.name, args.version)


if __name__ == "__main__":
    main()
